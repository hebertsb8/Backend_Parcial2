"""
Views para el sistema de predicción de ventas con Machine Learning.
"""
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAdminUser
from rest_framework.response import Response
from django.http import HttpResponse
from io import BytesIO
import openpyxl
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches

from sales.ml_predictor_simple import SimpleSalesPredictor
from sales.ml_model_manager import model_manager, get_predictor
try:
    from sales.ml_predictor_rf import RandomForestSalesPredictor
except Exception:  # pragma: no cover
    RandomForestSalesPredictor = None  # type: ignore
from sales.ml_data_generator import generate_sales_data
from sales.ml_auto_retrain import should_retrain_model, auto_retrain_if_needed, get_retrain_status


@api_view(['POST'])
@permission_classes([IsAdminUser])
def generate_demo_sales_data(request):
    """
    Genera datos sintéticos de ventas para demostración.
    
    POST /api/orders/ml/generate-demo-data/
    
    Body (opcional):
    {
        "clear_existing": true  // Elimina datos existentes antes de generar
    }
    
    Returns:
        Estadísticas de generación de datos
    """
    try:
        clear_existing = request.data.get('clear_existing', False)
        
        stats = generate_sales_data(clear_existing=clear_existing)
        
        return Response({
            'success': True,
            'message': 'Datos de demostración generados exitosamente',
            'data': stats
        }, status=status.HTTP_201_CREATED)
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
@permission_classes([IsAdminUser])
def train_model(request):
    """
    Entrena un nuevo modelo de predicción de ventas.
    
    POST /api/orders/ml/train/
    
    Body (opcional):
    {
        "notes": "Descripción del entrenamiento",
        "version": "v1.0"  // Opcional
    }
    
    Returns:
        Información del modelo entrenado
    """
    try:
        notes = request.data.get('notes', 'Entrenamiento manual desde API')
        version = request.data.get('version', None)
        algorithm = (request.data.get('algorithm') or 'linear').lower()
        
        # Entrenar modelo
        if algorithm == 'rf' and RandomForestSalesPredictor is not None:
            predictor = RandomForestSalesPredictor()
        else:
            predictor = SimpleSalesPredictor()
        metrics = predictor.train()
        
        # Guardar modelo
        model_info = model_manager.save_model(
            predictor,
            version=version,
            notes=notes,
            algorithm=algorithm
        )
        
        return Response({
            'success': True,
            'message': 'Modelo entrenado y guardado exitosamente',
            'model_info': model_info,
            'metrics': metrics,
            'algorithm': model_info.get('algorithm', algorithm)
        }, status=status.HTTP_201_CREATED)
        
    except ValueError as e:
        return Response({
            'success': False,
            'error': str(e),
            'hint': 'Asegúrate de tener al menos 30 días de datos de ventas. '
                   'Usa /api/orders/ml/generate-demo-data/ para generar datos de prueba.'
        }, status=status.HTTP_400_BAD_REQUEST)
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAdminUser])
def get_predictions(request):
    """
    Obtiene predicciones de ventas futuras.
    
    GET /api/orders/ml/predictions/
    
    Query params:
    - days: Número de días a predecir (default: 30, max: 365)
    
    Returns:
        Predicciones de ventas con intervalos de confianza
    """
    try:
        # Obtener parámetros
        days = int(request.query_params.get('days', 30))
        
        # Validar
        if days < 1 or days > 365:
            return Response({
                'success': False,
                'error': 'El parámetro "days" debe estar entre 1 y 365'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Obtener predictor
        predictor = get_predictor()
        
        # Generar predicciones
        predictions = predictor.predict(days=days)
        
        return Response({
            'success': True,
            'data': predictions
        })
        
    except ValueError as e:
        return Response({
            'success': False,
            'error': str(e),
            'hint': 'Entrena un modelo primero usando /api/orders/ml/train/'
        }, status=status.HTTP_400_BAD_REQUEST)
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAdminUser])
def get_forecast_components(request):
    """
    Obtiene información del modelo de predicción.
    
    GET /api/orders/ml/forecast-components/
    
    Returns:
        Información del modelo y características
    """
    try:
        predictor = get_predictor()
        
        components = {
            'model_type': 'Linear Regression with Polynomial Features',
            'features': [
                'Días desde inicio',
                'Día de la semana',
                'Fin de semana (sí/no)',
                'Estacionalidad mensual (seno/coseno)',
                'Estacionalidad semanal (seno/coseno)'
            ],
            'metrics': predictor.metrics,
            'training_period': {
                'start': predictor.metrics.get('start_date'),
                'end': predictor.metrics.get('end_date'),
                'days': predictor.metrics.get('training_samples')
            }
        }
        
        return Response({
            'success': True,
            'data': components
        })
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAdminUser])
def get_model_performance(request):
    """
    Obtiene métricas de rendimiento del modelo actual.
    
    GET /api/orders/ml/performance/
    
    Returns:
        Métricas de error del modelo
    """
    try:
        predictor = get_predictor()
        performance = predictor.get_historical_performance()
        
        return Response({
            'success': True,
            'data': performance
        })
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAdminUser])
def list_models(request):
    """
    Lista todos los modelos guardados.
    
    GET /api/orders/ml/models/
    
    Returns:
        Lista de modelos con su información
    """
    try:
        models = model_manager.list_models()
        current_model = model_manager.get_current_model_info()
        
        return Response({
            'success': True,
            'data': {
                'models': models,
                'current_model': current_model
            }
        })
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAdminUser])
def set_current_model(request):
    """
    Establece qué modelo usar como actual.
    
    POST /api/orders/ml/models/set-current/
    
    Body:
    {
        "version": "20241020_153045"
    }
    
    Returns:
        Información del modelo establecido
    """
    try:
        version = request.data.get('version')
        
        if not version:
            return Response({
                'success': False,
                'error': 'Se requiere el parámetro "version"'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        model_info = model_manager.set_current_model(version)
        
        return Response({
            'success': True,
            'message': f'Modelo {version} establecido como actual',
            'data': model_info
        })
        
    except ValueError as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_404_NOT_FOUND)
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['DELETE'])
@permission_classes([IsAdminUser])
def delete_model(request, version):
    """
    Elimina un modelo guardado.
    
    DELETE /api/orders/ml/models/{version}/
    
    Returns:
        Confirmación de eliminación
    """
    try:
        model_manager.delete_model(version)
        
        return Response({
            'success': True,
            'message': f'Modelo {version} eliminado exitosamente'
        })
        
    except ValueError as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_400_BAD_REQUEST)
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAdminUser])
def check_retrain_status(request):
    """
    Verifica si el modelo necesita ser reentrenado.

    GET /api/orders/ml/retrain/status/

    Returns:
        Estado del sistema de reentrenamiento automático
    """
    try:
        info = get_retrain_status()

        return Response({
            'success': True,
            'data': info
        })

    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAdminUser])
def auto_retrain(request):
    """
    Reentrena el modelo automáticamente si es necesario.

    POST /api/orders/ml/retrain/auto/

    Body (opcional):
    {
        "force": false  // Forzar reentrenamiento
    }

    Returns:
        Resultado del reentrenamiento
    """
    try:
        force = request.data.get('force', False)

        result = auto_retrain_if_needed(force=force)

        if result['retrained']:
            return Response({
                'success': True,
                'message': 'Modelo reentrenado exitosamente',
                'data': result
            }, status=status.HTTP_201_CREATED)
        elif result['error']:
            return Response({
                'success': False,
                'error': result['error']
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        else:
            return Response({
                'success': True,
                'message': 'No fue necesario reentrenar',
                'data': result
            })

    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAdminUser])
def ml_dashboard(request):
    """
    Dashboard completo del sistema ML.
    
    GET /api/orders/ml/dashboard/?format=json|pdf|excel|word
    
    Returns:
        Resumen completo: modelo actual, predicciones, performance
    """
    try:
        # Obtener modelo actual
        current_model = model_manager.get_current_model_info()
        
        if current_model is None:
            return Response({
                'success': False,
                'error': 'No hay modelo entrenado',
                'hint': 'Entrena un modelo usando /api/orders/ml/train/'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Obtener predictor
        predictor = get_predictor()
        
        # Generar predicciones a 30 días
        predictions = predictor.predict(days=30)
        
        # Obtener performance
        performance = predictor.get_historical_performance()
        
        # Resumen
        dashboard = {
            'current_model': current_model,
            'predictions_30_days': {
                'total_predicted': predictions['summary']['total_predicted_sales'],
                'average_daily': predictions['summary']['average_daily_sales'],
                'growth_rate': predictions['summary']['growth_rate_percent']
            },
            'performance': performance,
            'last_updated': current_model['saved_at']
        }
        
        # Verificar formato
        format_type = request.GET.get('format', 'json').lower()
        
        if format_type == 'json':
            return Response({
                'success': True,
                'data': dashboard
            })
        elif format_type == 'pdf':
            return _generate_ml_dashboard_pdf(dashboard)
        elif format_type == 'excel':
            return _generate_ml_dashboard_excel(dashboard)
        elif format_type == 'word':
            return _generate_ml_dashboard_word(dashboard)
        else:
            return Response(
                {'error': 'Formato no soportado. Use: json, pdf, excel, word'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def _generate_ml_dashboard_pdf(data):
    """Genera PDF del dashboard ML"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Título
    elements.append(Paragraph("Dashboard de Machine Learning", styles['Heading1']))
    elements.append(Spacer(1, 12))

    # Modelo actual
    if 'current_model' in data:
        elements.append(Paragraph("Modelo Actual", styles['Heading2']))
        model_info = data['current_model']
        elements.append(Paragraph(f"Nombre: {model_info.get('name', 'N/A')}", styles['Normal']))
        elements.append(Paragraph(f"Tipo: {model_info.get('type', 'N/A')}", styles['Normal']))
        elements.append(Paragraph(f"Precisión: {model_info.get('accuracy', 'N/A')}", styles['Normal']))
        elements.append(Spacer(1, 12))

    # Predicciones
    if 'predictions_30_days' in data:
        elements.append(Paragraph("Predicciones a 30 días", styles['Heading2']))
        pred = data['predictions_30_days']
        elements.append(Paragraph(f"Total predicho: {pred.get('total_predicted', 'N/A')}", styles['Normal']))
        elements.append(Paragraph(f"Promedio diario: {pred.get('average_daily', 'N/A')}", styles['Normal']))
        elements.append(Paragraph(f"Tasa de crecimiento: {pred.get('growth_rate', 'N/A')}%", styles['Normal']))
        elements.append(Spacer(1, 12))

    # Performance
    if 'performance' in data and isinstance(data['performance'], dict):
        elements.append(Paragraph("Performance Histórico", styles['Heading2']))
        perf = data['performance']
        table_data = [['Métrica', 'Valor']]
        for key, value in perf.items():
            table_data.append([str(key), str(value)])
        
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(table)

    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="dashboard_ml.pdf"'
    return response


def _generate_ml_dashboard_excel(data):
    """Genera Excel del dashboard ML"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dashboard ML"  # type: ignore[reportOptionalMemberAccess]

    row = 1
    # Título
    ws.cell(row=row, column=1, value="Dashboard de Machine Learning")  # type: ignore[reportOptionalMemberAccess]
    row += 2

    # Modelo actual
    if 'current_model' in data:
        ws.cell(row=row, column=1, value="Modelo Actual")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        model_info = data['current_model']
        ws.cell(row=row, column=1, value=f"Nombre: {model_info.get('name', 'N/A')}")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        ws.cell(row=row, column=1, value=f"Tipo: {model_info.get('type', 'N/A')}")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        ws.cell(row=row, column=1, value=f"Precisión: {model_info.get('accuracy', 'N/A')}")  # type: ignore[reportOptionalMemberAccess]
        row += 2

    # Predicciones
    if 'predictions_30_days' in data:
        ws.cell(row=row, column=1, value="Predicciones a 30 días")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        pred = data['predictions_30_days']
        ws.cell(row=row, column=1, value=f"Total predicho: {pred.get('total_predicted', 'N/A')}")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        ws.cell(row=row, column=1, value=f"Promedio diario: {pred.get('average_daily', 'N/A')}")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        ws.cell(row=row, column=1, value=f"Tasa de crecimiento: {pred.get('growth_rate', 'N/A')}%")  # type: ignore[reportOptionalMemberAccess]
        row += 2

    # Performance
    if 'performance' in data and isinstance(data['performance'], dict):
        ws.cell(row=row, column=1, value="Performance Histórico")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        ws.cell(row=row, column=2, value="Valor")  # type: ignore[reportOptionalMemberAccess]
        row += 1
        for key, value in data['performance'].items():
            ws.cell(row=row, column=1, value=str(key))  # type: ignore[reportOptionalMemberAccess]
            ws.cell(row=row, column=2, value=str(value))  # type: ignore[reportOptionalMemberAccess]
            row += 1

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="dashboard_ml.xlsx"'
    return response


def _generate_ml_dashboard_word(data):
    """Genera Word del dashboard ML"""
    doc = Document()
    doc.add_heading('Dashboard de Machine Learning', 0)

    # Modelo actual
    if 'current_model' in data:
        doc.add_heading('Modelo Actual', level=1)
        model_info = data['current_model']
        doc.add_paragraph(f"Nombre: {model_info.get('name', 'N/A')}")
        doc.add_paragraph(f"Tipo: {model_info.get('type', 'N/A')}")
        doc.add_paragraph(f"Precisión: {model_info.get('accuracy', 'N/A')}")

    # Predicciones
    if 'predictions_30_days' in data:
        doc.add_heading('Predicciones a 30 días', level=1)
        pred = data['predictions_30_days']
        doc.add_paragraph(f"Total predicho: {pred.get('total_predicted', 'N/A')}")
        doc.add_paragraph(f"Promedio diario: {pred.get('average_daily', 'N/A')}")
        doc.add_paragraph(f"Tasa de crecimiento: {pred.get('growth_rate', 'N/A')}%")

    # Performance
    if 'performance' in data and isinstance(data['performance'], dict):
        doc.add_heading('Performance Histórico', level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells  # type: ignore[reportOptionalMemberAccess]
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        for key, value in data['performance'].items():
            row_cells = table.add_row().cells  # type: ignore[reportOptionalMemberAccess]
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="dashboard_ml.docx"'
    return response
