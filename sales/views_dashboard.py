# sales/views_dashboard.py
"""
Vistas para el dashboard en tiempo real
"""

from rest_framework import views, status
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from api.permissions import IsAdminUser
from .analytics import DashboardAnalytics
from django.core.cache import cache
from django.http import HttpResponse
from io import BytesIO
import openpyxl
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches
import json


class RealTimeDashboardView(views.APIView):
    """
    GET /api/orders/dashboard/realtime/
    
    Dashboard en tiempo real con estadísticas actualizadas.
    Los datos se cachean por 5 minutos para mejorar el rendimiento.
    """
    permission_classes = [IsAdminUser]
    
    def get(self, request):
        """Obtiene estadísticas en tiempo real"""
        
        # Intentar obtener del cache primero
        cache_key = 'dashboard_realtime_stats'
        cached_data = cache.get(cache_key)
        
        if cached_data and not request.GET.get('format'):
            cached_data['from_cache'] = True
            return Response(cached_data)
        
        try:
            # Generar estadísticas frescas
            stats = DashboardAnalytics.get_real_time_stats()
            stats['from_cache'] = False
            
            # Verificar si se solicita descarga
            format_type = request.GET.get('format', 'json').lower()
            
            if format_type == 'json':
                # Cachear por 5 minutos solo para JSON
                if not cached_data:
                    cache.set(cache_key, stats, 300)
                return Response(stats, status=status.HTTP_200_OK)
            elif format_type == 'pdf':
                return generate_dashboard_pdf(stats, 'Dashboard en Tiempo Real')
            elif format_type == 'excel':
                return generate_dashboard_excel(stats, 'Dashboard en Tiempo Real')
            elif format_type == 'word':
                return generate_dashboard_word(stats, 'Dashboard en Tiempo Real')
            else:
                return Response(
                    {'error': 'Formato no soportado. Use: json, pdf, excel, word'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        except Exception as e:
            return Response(
                {'error': f'Error al generar estadísticas: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class ProductPerformanceView(views.APIView):
    """
    GET /api/orders/dashboard/products/
    GET /api/orders/dashboard/products/{product_id}/
    
    Análisis de rendimiento de productos
    """
    permission_classes = [IsAdminUser]
    
    def get(self, request, product_id=None):
        """Obtiene análisis de rendimiento de productos"""
        
        try:
            performance = DashboardAnalytics.get_product_performance(product_id)
            data = {
                'count': len(performance),
                'products': performance
            }
            
            # Verificar si se solicita descarga
            format_type = request.GET.get('format', 'json').lower()
            
            if format_type == 'json':
                return Response(data, status=status.HTTP_200_OK)
            elif format_type == 'pdf':
                return generate_dashboard_pdf(data, f'Rendimiento de Productos{" - " + str(product_id) if product_id else ""}')
            elif format_type == 'excel':
                return generate_dashboard_excel(data, f'Rendimiento de Productos{" - " + str(product_id) if product_id else ""}')
            elif format_type == 'word':
                return generate_dashboard_word(data, f'Rendimiento de Productos{" - " + str(product_id) if product_id else ""}')
            else:
                return Response(
                    {'error': 'Formato no soportado. Use: json, pdf, excel, word'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class CustomerInsightsView(views.APIView):
    """
    GET /api/orders/dashboard/customers/
    GET /api/orders/dashboard/customers/{customer_id}/
    
    Insights de clientes y análisis de comportamiento
    """
    permission_classes = [IsAdminUser]
    
    def get(self, request, customer_id=None):
        """Obtiene insights de clientes"""
        
        try:
            insights = DashboardAnalytics.get_customer_insights(customer_id)
            data = {
                'count': len(insights),
                'customers': insights
            }
            
            # Verificar si se solicita descarga
            format_type = request.GET.get('format', 'json').lower()
            
            if format_type == 'json':
                return Response(data, status=status.HTTP_200_OK)
            elif format_type == 'pdf':
                return generate_dashboard_pdf(data, f'Insights de Clientes{" - " + str(customer_id) if customer_id else ""}')
            elif format_type == 'excel':
                return generate_dashboard_excel(data, f'Insights de Clientes{" - " + str(customer_id) if customer_id else ""}')
            elif format_type == 'word':
                return generate_dashboard_word(data, f'Insights de Clientes{" - " + str(customer_id) if customer_id else ""}')
            else:
                return Response(
                    {'error': 'Formato no soportado. Use: json, pdf, excel, word'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class InvalidateCacheView(views.APIView):
    """
    POST /api/orders/dashboard/invalidate-cache/
    
    Invalida el cache del dashboard para forzar actualización
    """
    permission_classes = [IsAdminUser]
    
    def post(self, request):
        """Invalida el cache"""
        
        cache.delete('dashboard_realtime_stats')
        
        return Response({
            'message': 'Cache invalidado exitosamente'
        }, status=status.HTTP_200_OK)


# Funciones globales para generar archivos de descarga
def generate_dashboard_pdf(data, title):
    """Genera PDF del dashboard"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Título
    elements.append(Paragraph(title, styles['Heading1']))
    elements.append(Spacer(1, 12))

    # Timestamp
    if 'timestamp' in data:
        elements.append(Paragraph(f"Generado: {data['timestamp']}", styles['Normal']))
        elements.append(Spacer(1, 12))

    # Procesar datos
    for section, content in data.items():
        if section in ['timestamp', 'from_cache']:
            continue
        
        elements.append(Paragraph(f"{section.upper()}", styles['Heading2']))
        
        if isinstance(content, dict):
            # Crear tabla para diccionarios
            table_data = [['Métrica', 'Valor']]
            for key, value in content.items():
                table_data.append([str(key), str(value)])
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(table)
        
        elif isinstance(content, list) and content:
            # Crear tabla para listas
            if isinstance(content[0], dict):
                headers = list(content[0].keys())
                table_data = [headers] + [[str(item.get(h, '')) for h in headers] for item in content[:20]]
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(table)
        
        elements.append(Spacer(1, 12))

    doc.build(elements)
    
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{title.replace(" ", "_")}.pdf"'
    return response


def generate_dashboard_excel(data, title):
    """Genera Excel del dashboard"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dashboard"  # type: ignore[reportOptionalMemberAccess]

    row = 1
    # Título
    ws.cell(row=row, column=1, value=title)  # type: ignore[reportOptionalMemberAccess]
    row += 2

    # Timestamp
    if 'timestamp' in data:
        ws.cell(row=row, column=1, value=f"Generado: {data['timestamp']}")  # type: ignore[reportOptionalMemberAccess]
        row += 2

    # Procesar datos
    for section, content in data.items():
        if section in ['timestamp', 'from_cache']:
            continue
        
        ws.cell(row=row, column=1, value=section.upper())  # type: ignore[reportOptionalMemberAccess]
        row += 1
        
        if isinstance(content, dict):
            # Headers
            ws.cell(row=row, column=1, value="Métrica")  # type: ignore[reportOptionalMemberAccess]
            ws.cell(row=row, column=2, value="Valor")  # type: ignore[reportOptionalMemberAccess]
            row += 1
            
            for key, value in content.items():
                ws.cell(row=row, column=1, value=str(key))  # type: ignore[reportOptionalMemberAccess]
                ws.cell(row=row, column=2, value=str(value))  # type: ignore[reportOptionalMemberAccess]
                row += 1
        
        elif isinstance(content, list) and content:
            if isinstance(content[0], dict):
                headers = list(content[0].keys())
                for col, header in enumerate(headers, 1):
                    ws.cell(row=row, column=col, value=header)  # type: ignore[reportOptionalMemberAccess]
                row += 1
                
                for item in content[:50]:  # Limitar a 50 filas
                    for col, header in enumerate(headers, 1):
                        ws.cell(row=row, column=col, value=str(item.get(header, '')))  # type: ignore[reportOptionalMemberAccess]
                    row += 1
        
        row += 2  # Espacio entre secciones

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{title.replace(" ", "_")}.xlsx"'
    return response


def generate_dashboard_word(data, title):
    """Genera Word del dashboard"""
    doc = Document()
    doc.add_heading(title, 0)

    # Timestamp
    if 'timestamp' in data:
        doc.add_paragraph(f"Generado: {data['timestamp']}")

    # Procesar datos
    for section, content in data.items():
        if section in ['timestamp', 'from_cache']:
            continue
        
        doc.add_heading(section.upper(), level=1)
        
        if isinstance(content, dict):
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Métrica'
            hdr_cells[1].text = 'Valor'
            
            for key, value in content.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
        
        elif isinstance(content, list) and content:
            if isinstance(content[0], dict):
                headers = list(content[0].keys())
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for item in content[:30]:  # Limitar a 30 filas
                    row_cells = table.add_row().cells
                    for i, header in enumerate(headers):
                        row_cells[i].text = str(item.get(header, ''))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{title.replace(" ", "_")}.docx"'
    return response
